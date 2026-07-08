"""Document parsers: file → plain text.

Markdown/text/HTML parse with the standard library alone; PDF and DOCX
use ``pypdfium2`` / ``python-docx``, imported lazily so the platform
never pays for parsers it does not use — a missing optional dependency
surfaces as a precise ``ConfigurationError`` at ingestion time, not an
ImportError at startup.
"""

from __future__ import annotations

from html.parser import HTMLParser
from pathlib import Path

from core.exceptions import ConfigurationError, DataSourceError
from core.logging import get_logger

logger = get_logger("rag.parsers")

#: Extension -> (media_type, parser function name). The registry the loader consults.
SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
    ".html": "html",
    ".htm": "html",
    ".pdf": "pdf",
    ".docx": "docx",
}


def parse_file(path: Path) -> tuple[str, str]:
    """Return ``(media_type, extracted_text)`` for a supported file."""
    media_type = SUPPORTED_EXTENSIONS.get(path.suffix.lower())
    if media_type is None:
        raise DataSourceError(
            f"Unsupported knowledge file type '{path.suffix}' ({path.name}). "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    try:
        if media_type in ("markdown", "text"):
            return media_type, path.read_text(encoding="utf-8")
        if media_type == "html":
            return media_type, _parse_html(path.read_text(encoding="utf-8"))
        if media_type == "pdf":
            return media_type, _parse_pdf(path)
        return media_type, _parse_docx(path)
    except (OSError, UnicodeDecodeError) as exc:
        raise DataSourceError(f"Cannot read knowledge file {path}: {exc}") from exc


class _TextExtractor(HTMLParser):
    """Minimal, dependency-free HTML → text (headings kept as markdown)."""

    _SKIP = {"script", "style", "noscript"}
    _HEADINGS = {"h1": "#", "h2": "##", "h3": "###", "h4": "####"}

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skipping = 0

    def handle_starttag(self, tag: str, attrs: list) -> None:
        if tag in self._SKIP:
            self._skipping += 1
        elif tag in self._HEADINGS:
            self._parts.append(f"\n\n{self._HEADINGS[tag]} ")
        elif tag in ("p", "div", "li", "br", "tr"):
            self._parts.append("\n\n" if tag in ("p", "div") else "\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP and self._skipping:
            self._skipping -= 1

    def handle_data(self, data: str) -> None:
        if not self._skipping and data.strip():
            self._parts.append(data.strip() + " ")

    def text(self) -> str:
        return "".join(self._parts).strip()


def _parse_html(markup: str) -> str:
    extractor = _TextExtractor()
    extractor.feed(markup)
    return extractor.text()


def _parse_pdf(path: Path) -> str:
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise ConfigurationError(
            "PDF ingestion requires 'pypdfium2': pip install pypdfium2"
        ) from exc
    document = pdfium.PdfDocument(str(path))
    try:
        pages = []
        for page in document:
            text_page = page.get_textpage()
            pages.append(text_page.get_text_bounded())
            text_page.close()
            page.close()
        return "\n\n".join(pages)
    finally:
        document.close()


def _parse_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ConfigurationError(
            "DOCX ingestion requires 'python-docx': pip install python-docx"
        ) from exc
    parsed = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in parsed.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "2"
            parts.append(f"{'#' * min(int(level), 4)} {text}")
        else:
            parts.append(text)
    return "\n\n".join(parts)
